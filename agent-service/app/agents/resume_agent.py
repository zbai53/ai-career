import json
import logging
import os
import re
import time
from pathlib import Path

import anthropic
import pdfplumber
from docx import Document
from pydantic import ValidationError

from app.models.resume import ParsedResume

logger = logging.getLogger(__name__)

_MODEL = "claude-haiku-4-5-20251001"
_MAX_TOKENS = 8192
_MIN_TEXT_CHARS = 50       # below this → scanned/image PDF
_MAX_TEXT_CHARS = 15_000   # truncate before sending to Claude

# JSON schema embedded once at import time so the prompt stays lean
_RESUME_SCHEMA = json.dumps(ParsedResume.model_json_schema(), indent=2)

_SYSTEM_PROMPT = f"""\
You are an expert resume parser. Given the raw text of a resume, extract all \
information and return ONLY a single valid JSON object that matches the schema \
below — no prose, no markdown fences, no extra keys.

SCHEMA:
{_RESUME_SCHEMA}

RULES:
1. Return only the JSON object, nothing else.
2. Dates must be in YYYY-MM format. If only a year is given use YYYY-01. \
   If the end date is "Present" / "Current", set end_date to null and \
   is_current to true.
3. If a section is absent from the resume, use an empty list [] or null as \
   appropriate per the schema.
4. When a candidate held multiple roles at the same company, emit a separate \
   experience entry for each role.
5. Set parse_confidence to a float between 0.0 and 1.0 that reflects how \
   completely the resume could be extracted: 0.9–1.0 for clean, complete \
   resumes; 0.6–0.9 for resumes with some ambiguity; below 0.6 for heavily \
   formatted or sparse documents.
6. For the raw_text field, include only the first 500 characters of the original \
   text followed by '...[truncated]'. Do NOT include the full raw text.\
"""

_RETRY_SUFFIX = """

IMPORTANT: Your previous response could not be validated against the schema. \
Common mistakes:
- Missing required fields: contact (object), raw_text (string), parse_confidence (float).
- Wrong date format — must be YYYY-MM or null.
- gpa must be a float between 0.0 and 4.0 or null.
- parse_confidence must be a float between 0.0 and 1.0.
Return ONLY the corrected JSON object.\
"""


class UnsupportedFileTypeError(Exception):
    pass


class TextExtractionError(Exception):
    pass


class ResumeParseError(Exception):
    pass


class ResumeAgent:
    def __init__(self) -> None:
        api_key = os.getenv("ANTHROPIC_API_KEY")
        if not api_key:
            raise EnvironmentError("ANTHROPIC_API_KEY is not set")
        self._client = anthropic.Anthropic(api_key=api_key)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def parse(self, file_path: str) -> ParsedResume:
        """Parse a resume file (.pdf or .docx) and return a validated ParsedResume."""
        path = Path(file_path)
        suffix = path.suffix.lower()

        t0 = time.perf_counter()

        if suffix == ".pdf":
            raw_text = self._extract_text_from_pdf(file_path)
        elif suffix in (".docx", ".doc"):
            raw_text = self._extract_text_from_docx(file_path)
        else:
            raise UnsupportedFileTypeError(
                f"Unsupported file type '{suffix}'. Supported: .pdf, .docx"
            )

        extraction_ms = (time.perf_counter() - t0) * 1000
        logger.info("Text extracted from %s in %.0f ms (%d chars)", path.name, extraction_ms, len(raw_text))

        # Truncate before sending to Claude to stay within token budget
        claude_text = raw_text
        if len(raw_text) > _MAX_TEXT_CHARS:
            logger.warning(
                "Resume text truncated from %d to %d chars for '%s'",
                len(raw_text), _MAX_TEXT_CHARS, path.name,
            )
            claude_text = raw_text[:_MAX_TEXT_CHARS]

        parsed = self._call_claude(claude_text, strict=False)
        # Always store the full extracted text on the result
        parsed.raw_text = raw_text

        total_ms = (time.perf_counter() - t0) * 1000
        logger.info("Resume parsed in %.0f ms total (confidence=%.2f)", total_ms, parsed.parse_confidence)
        return parsed

    # ------------------------------------------------------------------
    # Text extraction helpers
    # ------------------------------------------------------------------

    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF, page by page, to handle multi-column layouts."""
        try:
            pages: list[str] = []
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text(x_tolerance=3, y_tolerance=3)
                    if text:
                        pages.append(text.strip())
                    else:
                        logger.debug("Page %d yielded no text (possibly image-based)", i + 1)

            if not pages:
                raise TextExtractionError(
                    "PDF appears to be scanned/image-based. Text extraction failed."
                )

            joined = "\n\n".join(pages)

            if len(joined.strip()) < _MIN_TEXT_CHARS:
                raise TextExtractionError(
                    "PDF appears to be scanned/image-based. Text extraction failed."
                )

            return self._clean_pdf_text(joined)
        except TextExtractionError:
            raise
        except Exception as exc:
            raise TextExtractionError(f"Failed to read PDF '{file_path}': {exc}") from exc

    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX, including paragraphs and table cells."""
        try:
            doc = Document(file_path)
            parts: list[str] = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        parts.append("  ".join(row_cells))

            if not parts:
                raise TextExtractionError(f"No text could be extracted from DOCX: {file_path}")
            return "\n".join(parts)
        except TextExtractionError:
            raise
        except Exception as exc:
            raise TextExtractionError(f"Failed to read DOCX '{file_path}': {exc}") from exc

    def _clean_pdf_text(self, text: str) -> str:
        """Remove common PDF artifacts: repeated whitespace, lone page numbers, header/footer noise."""
        # Collapse runs of 3+ blank lines to a single blank line
        text = re.sub(r"\n{3,}", "\n\n", text)
        # Remove lines that are solely a page number (digits, optionally preceded by "Page")
        text = re.sub(r"(?im)^(page\s*)?\d+\s*$", "", text)
        # Collapse multiple spaces/tabs to a single space on each line
        lines = [re.sub(r"[ \t]{2,}", " ", line) for line in text.splitlines()]
        # Drop lines that are pure whitespace after cleanup
        lines = [line for line in lines if line.strip()]
        return "\n".join(lines)

    # ------------------------------------------------------------------
    # Claude interaction
    # ------------------------------------------------------------------

    def _call_claude(self, raw_text: str, strict: bool) -> ParsedResume:
        user_content = f"Parse the following resume:\n\n{raw_text}"
        if strict:
            user_content += _RETRY_SUFFIX

        t0 = time.perf_counter()
        response = self._client.messages.create(
            model=_MODEL,
            max_tokens=_MAX_TOKENS,
            system=_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": user_content}],
        )
        elapsed_ms = (time.perf_counter() - t0) * 1000

        usage = response.usage
        logger.info(
            "Claude responded in %.0f ms | input_tokens=%d output_tokens=%d",
            elapsed_ms,
            usage.input_tokens,
            usage.output_tokens,
        )

        raw_json = response.content[0].text.strip()

        # Strip accidental markdown fences that some models emit despite instructions
        if raw_json.startswith("```"):
            lines = raw_json.splitlines()
            raw_json = "\n".join(
                line for line in lines if not line.startswith("```")
            ).strip()

        try:
            data = json.loads(raw_json)
            # Ensure raw_text is the actual extracted text, not a truncated summary
            data["raw_text"] = raw_text
            return ParsedResume.model_validate(data)
        except (json.JSONDecodeError, ValidationError) as exc:
            if strict:
                raise ResumeParseError(
                    f"Claude response failed validation after retry: {exc}\n\nResponse:\n{raw_json}"
                ) from exc
            logger.warning("Initial parse failed (%s), retrying with stricter prompt", exc)
            return self._call_claude(raw_text, strict=True)

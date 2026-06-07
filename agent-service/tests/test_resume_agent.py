"""
Tests for ResumeAgent.

Real Anthropic API calls are never made — all tests mock self._client.messages.create.
Temp files created by fixtures are cleaned up automatically by pytest's tmp_path.
"""

import json
import os
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas

from app.agents.resume_agent import (
    ResumeAgent,
    ResumeParseError,
    TextExtractionError,
    UnsupportedFileTypeError,
)
from app.models.resume import ParsedResume

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_VALID_RESUME_PAYLOAD = {
    "contact": {
        "name": "Jane Doe",
        "email": "jane@example.com",
        "phone": "555-1234",
        "address": "Austin, TX",
        "linkedin_url": None,
        "github_url": None,
        "portfolio_url": None,
    },
    "summary": "Experienced software engineer.",
    "education": [
        {
            "institution": "UT Austin",
            "degree": "Bachelor of Science",
            "field_of_study": "Computer Science",
            "start_date": "2014-08",
            "end_date": "2018-05",
            "gpa": 3.7,
            "highlights": [],
        }
    ],
    "experience": [
        {
            "company": "Acme Corp",
            "title": "Software Engineer",
            "location": "Austin, TX",
            "start_date": "2018-06",
            "end_date": None,
            "is_current": True,
            "bullets": ["Built core API serving 1M requests/day"],
            "technologies": ["Python", "PostgreSQL"],
        }
    ],
    "projects": [],
    "certifications": [],
    "skills": [
        {"name": "Python", "category": "language", "proficiency": "expert"},
    ],
    "raw_text": "Jane Doe\njane@example.com",
    "parse_confidence": 0.92,
}


def _make_fake_response(text: str) -> MagicMock:
    """Build a minimal object that mirrors anthropic.types.Message."""
    content_block = SimpleNamespace(text=text)
    usage = SimpleNamespace(input_tokens=100, output_tokens=200)
    return SimpleNamespace(content=[content_block], usage=usage)


def _make_agent() -> ResumeAgent:
    """Instantiate ResumeAgent with a dummy API key (no real network calls)."""
    with patch.dict(os.environ, {"ANTHROPIC_API_KEY": "test-key"}):
        with patch("anthropic.Anthropic"):
            return ResumeAgent()


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def pdf_file(tmp_path: Path) -> Path:
    """Create a single-page PDF containing known text."""
    path = tmp_path / "test_resume.pdf"
    c = canvas.Canvas(str(path), pagesize=LETTER)
    c.setFont("Helvetica", 12)
    c.drawString(72, 720, "Jane Doe")
    c.drawString(72, 700, "jane@example.com  |  555-1234")
    c.drawString(72, 680, "Software Engineer with 5 years of experience.")
    c.save()
    return path


@pytest.fixture
def docx_file(tmp_path: Path) -> Path:
    """Create a DOCX with paragraphs and a simple table containing known text."""
    path = tmp_path / "test_resume.docx"
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("jane@example.com  |  555-1234")
    doc.add_paragraph("Software Engineer with 5 years of experience.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Python"
    table.cell(0, 1).text = "expert"
    table.cell(1, 0).text = "PostgreSQL"
    table.cell(1, 1).text = "proficient"
    doc.save(str(path))
    return path


@pytest.fixture
def txt_file(tmp_path: Path) -> Path:
    path = tmp_path / "resume.txt"
    path.write_text("Jane Doe\njane@example.com")
    return path


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


class TestExtractTextFromPdf:
    def test_returns_expected_content(self, pdf_file: Path) -> None:
        agent = _make_agent()
        text = agent._extract_text_from_pdf(str(pdf_file))

        assert "Jane Doe" in text
        assert "jane@example.com" in text
        assert "Software Engineer" in text

    def test_raises_on_nonexistent_file(self, tmp_path: Path) -> None:
        agent = _make_agent()
        with pytest.raises(TextExtractionError):
            agent._extract_text_from_pdf(str(tmp_path / "missing.pdf"))


class TestExtractTextFromDocx:
    def test_returns_paragraph_content(self, docx_file: Path) -> None:
        agent = _make_agent()
        text = agent._extract_text_from_docx(str(docx_file))

        assert "Jane Doe" in text
        assert "jane@example.com" in text
        assert "Software Engineer" in text

    def test_returns_table_content(self, docx_file: Path) -> None:
        agent = _make_agent()
        text = agent._extract_text_from_docx(str(docx_file))

        # Table cells must be included
        assert "Python" in text
        assert "PostgreSQL" in text

    def test_raises_on_nonexistent_file(self, tmp_path: Path) -> None:
        agent = _make_agent()
        with pytest.raises(TextExtractionError):
            agent._extract_text_from_docx(str(tmp_path / "missing.docx"))


class TestUnsupportedFileType:
    def test_txt_raises_unsupported_error(self, txt_file: Path) -> None:
        agent = _make_agent()
        with pytest.raises(UnsupportedFileTypeError) as exc_info:
            agent.parse(str(txt_file))
        assert ".txt" in str(exc_info.value)

    def test_error_message_lists_supported_types(self, txt_file: Path) -> None:
        agent = _make_agent()
        with pytest.raises(UnsupportedFileTypeError) as exc_info:
            agent.parse(str(txt_file))
        msg = str(exc_info.value).lower()
        assert "pdf" in msg
        assert "docx" in msg


class TestParseReturnsValidSchema:
    def test_returns_parsed_resume_instance(self, pdf_file: Path) -> None:
        agent = _make_agent()
        payload = dict(_VALID_RESUME_PAYLOAD)
        agent._client.messages.create.return_value = _make_fake_response(
            json.dumps(payload)
        )

        result = agent.parse(str(pdf_file))

        assert isinstance(result, ParsedResume)

    def test_contact_fields_populated(self, pdf_file: Path) -> None:
        agent = _make_agent()
        agent._client.messages.create.return_value = _make_fake_response(
            json.dumps(_VALID_RESUME_PAYLOAD)
        )

        result = agent.parse(str(pdf_file))

        assert result.contact.name == "Jane Doe"
        assert result.contact.email == "jane@example.com"

    def test_raw_text_is_extracted_content_not_claude_summary(self, pdf_file: Path) -> None:
        """raw_text must come from local extraction, not from Claude's response."""
        agent = _make_agent()
        payload = dict(_VALID_RESUME_PAYLOAD)
        payload["raw_text"] = "Claude truncated summary"  # agent should override this
        agent._client.messages.create.return_value = _make_fake_response(
            json.dumps(payload)
        )

        result = agent.parse(str(pdf_file))

        # Should contain actual PDF text, not Claude's placeholder
        assert "Jane Doe" in result.raw_text

    def test_parse_confidence_within_bounds(self, pdf_file: Path) -> None:
        agent = _make_agent()
        agent._client.messages.create.return_value = _make_fake_response(
            json.dumps(_VALID_RESUME_PAYLOAD)
        )

        result = agent.parse(str(pdf_file))

        assert 0.0 <= result.parse_confidence <= 1.0

    def test_experience_list_populated(self, pdf_file: Path) -> None:
        agent = _make_agent()
        agent._client.messages.create.return_value = _make_fake_response(
            json.dumps(_VALID_RESUME_PAYLOAD)
        )

        result = agent.parse(str(pdf_file))

        assert len(result.experience) == 1
        assert result.experience[0].company == "Acme Corp"
        assert result.experience[0].is_current is True

    def test_markdown_fences_stripped(self, pdf_file: Path) -> None:
        """Agent must handle Claude responses wrapped in ```json fences."""
        agent = _make_agent()
        fenced = f"```json\n{json.dumps(_VALID_RESUME_PAYLOAD)}\n```"
        agent._client.messages.create.return_value = _make_fake_response(fenced)

        result = agent.parse(str(pdf_file))

        assert isinstance(result, ParsedResume)
        assert result.contact.name == "Jane Doe"


class TestParseRetryOnInvalidJson:
    def test_retries_and_succeeds_on_second_call(self, pdf_file: Path) -> None:
        agent = _make_agent()
        valid_response = _make_fake_response(json.dumps(_VALID_RESUME_PAYLOAD))
        agent._client.messages.create.side_effect = [
            _make_fake_response("this is not valid json {{{"),
            valid_response,
        ]

        result = agent.parse(str(pdf_file))

        assert isinstance(result, ParsedResume)
        assert agent._client.messages.create.call_count == 2

    def test_retry_call_contains_stricter_prompt(self, pdf_file: Path) -> None:
        agent = _make_agent()
        valid_response = _make_fake_response(json.dumps(_VALID_RESUME_PAYLOAD))
        agent._client.messages.create.side_effect = [
            _make_fake_response("not json"),
            valid_response,
        ]

        agent.parse(str(pdf_file))

        _, retry_kwargs = agent._client.messages.create.call_args_list[1]
        retry_messages = retry_kwargs.get("messages") or agent._client.messages.create.call_args_list[1][0][0] if agent._client.messages.create.call_args_list[1][0] else []
        # Second call's user message must include the retry suffix
        second_call = agent._client.messages.create.call_args_list[1]
        user_msg = second_call.kwargs.get("messages", second_call.args[0] if second_call.args else [])
        if not isinstance(user_msg, list):
            user_msg = second_call[1].get("messages", [])
        assert any("IMPORTANT" in m["content"] for m in user_msg if isinstance(m, dict))

    def test_raises_resume_parse_error_after_two_failures(self, pdf_file: Path) -> None:
        agent = _make_agent()
        agent._client.messages.create.side_effect = [
            _make_fake_response("not json at all"),
            _make_fake_response("still not json"),
        ]

        with pytest.raises(ResumeParseError):
            agent.parse(str(pdf_file))

        assert agent._client.messages.create.call_count == 2

    def test_invalid_schema_triggers_retry(self, pdf_file: Path) -> None:
        """A structurally valid JSON that fails Pydantic validation should also retry."""
        agent = _make_agent()
        bad_payload = {"contact": {}, "parse_confidence": 99.0}  # confidence out of range
        valid_response = _make_fake_response(json.dumps(_VALID_RESUME_PAYLOAD))
        agent._client.messages.create.side_effect = [
            _make_fake_response(json.dumps(bad_payload)),
            valid_response,
        ]

        result = agent.parse(str(pdf_file))

        assert isinstance(result, ParsedResume)
        assert agent._client.messages.create.call_count == 2


class TestEdgeCases:
    def test_scanned_pdf_raises_error(self, tmp_path: Path) -> None:
        """A PDF with no extractable text (scanned/image-based) must raise TextExtractionError."""
        # Build a PDF that pdfplumber opens successfully but yields no text
        path = tmp_path / "scanned.pdf"
        c = canvas.Canvas(str(path), pagesize=LETTER)
        # Save an entirely blank page — no drawString calls
        c.save()

        agent = _make_agent()
        with pytest.raises(TextExtractionError) as exc_info:
            agent._extract_text_from_pdf(str(path))

        msg = str(exc_info.value).lower()
        assert "scanned" in msg or "image" in msg or "text extraction failed" in msg

    def test_long_resume_truncated_before_claude(self, tmp_path: Path) -> None:
        """Text longer than 15 000 chars must be truncated before being sent to Claude."""
        # Create a PDF with known content; we'll mock _extract_text_from_pdf to return long text
        pdf_path = tmp_path / "long_resume.pdf"
        c = canvas.Canvas(str(pdf_path), pagesize=LETTER)
        c.drawString(72, 720, "Jane Doe")
        c.save()

        long_text = "A" * 20_000  # 20 000 chars — well above the 15 000 limit

        agent = _make_agent()
        agent._client.messages.create.return_value = _make_fake_response(
            json.dumps(_VALID_RESUME_PAYLOAD)
        )

        with patch.object(agent, "_extract_text_from_pdf", return_value=long_text):
            result = agent.parse(str(pdf_path))

        # Claude must have received at most 15 000 chars in the user message
        call_args = agent._client.messages.create.call_args
        user_msg = call_args.kwargs["messages"][0]["content"]
        # The user message is "Parse the following resume:\n\n" + text
        sent_text = user_msg.split("\n\n", 1)[1]
        assert len(sent_text) <= 15_000

        # But the stored raw_text on the result must be the full text
        assert result.raw_text == long_text
